"""
Text, Markdown, and DOCX Document Loader.
"""

from typing import List
from pathlib import Path
from langchain_community.document_loaders import TextLoader
from langchain_text_splitters import CharacterTextSplitter
from langchain_core.documents import Document


def load_text_document(
    file_path: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 100,
) -> List[Document]:
    """
    Loads a text, markdown, or docx file and splits it into chunked documents.
    """
    suffix = Path(file_path).suffix.lower()

    if suffix == ".docx":
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            docs = [Document(page_content=full_text, metadata={"source": file_path})]
        except Exception:
            loader = TextLoader(file_path, encoding="utf-8")
            docs = loader.load()
    else:
        loader = TextLoader(file_path, encoding="utf-8")
        docs = loader.load()

    # Character splitter preserving original logic
    splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )
    chunks = splitter.split_documents(docs)

    cleaned_chunks = []
    for chunk in chunks:
        chunk.page_content = chunk.page_content.strip()
        if chunk.page_content:
            cleaned_chunks.append(chunk)

    return cleaned_chunks
